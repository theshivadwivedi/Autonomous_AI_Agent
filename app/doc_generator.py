import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUTPUT_DIR = Path("app/output")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


def _add_markdown_text(paragraph, text: str) -> None:
    """
    Adds text to a paragraph, turning **bold** markers into actual
    bold runs instead of leaving literal asterisks in the document.
    """
    pos = 0
    for match in BOLD_PATTERN.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        run = paragraph.add_run(match.group(1))
        run.bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _is_standalone_bold_heading(line: str) -> bool:
    # e.g. "**Executive Summary:**" on its own line - the LLM is using
    # this as a sub-heading, so render it as one instead of bold text
    return bool(re.fullmatch(r"\*\*[^*]{2,60}:?\*\*", line))


def _add_section_content(document: Document, raw_content: str) -> None:
    """
    The LLM returns markdown-ish text (bold markers, numbered lists,
    bullet lists) inside what is otherwise plain content. python-docx
    doesn't render markdown, so without this the document shows literal
    ** characters. This walks the content line by line and maps each
    line to real Word formatting instead.
    """
    for raw_line in raw_content.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        if _is_standalone_bold_heading(line):
            document.add_heading(line.strip("*").strip(":").strip(), level=4)
            continue

        numbered = re.match(r"^\d+\.\s+(.*)", line)
        bulleted = re.match(r"^[\*\-]\s+(.*)", line)

        if numbered:
            # Word's "List Number" style shares one counter across the
            # whole document - multiple lists end up continuing each
            # other's numbering (1,2 then 3,4,5...) instead of each
            # restarting at 1. Using bullets sidesteps that entirely;
            # reading order still conveys the sequence.
            p = document.add_paragraph(style="List Bullet")
            _add_markdown_text(p, numbered.group(1))
        elif bulleted:
            p = document.add_paragraph(style="List Bullet")
            _add_markdown_text(p, bulleted.group(1))
        else:
            p = document.add_paragraph()
            _add_markdown_text(p, line)


def _make_title(request: str) -> str:
    cleaned = " ".join(request.split())
    if not cleaned:
        return "Business Document"
    if len(cleaned) > 70:
        cleaned = cleaned[:67].rsplit(" ", 1)[0] + "..."
    return cleaned[0].upper() + cleaned[1:]


def create_word_document(
    request: str,
    outputs: dict[str, str],
    tasks: list | None = None,
) -> str:
    """
    Create a professional Word document from the agent's plan and the
    completed task outputs.

    Args:
        request: Original user request.
        outputs: Dictionary of task title -> generated content.
        tasks: The plan's task list (optional) - when provided, the
               document shows the agent's plan before the content it
               produced, so it's visible that planning and execution
               were separate steps.

    Returns:
        Path to the generated .docx file.
    """

    document = Document()

    title = document.add_heading(_make_title(request), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Generated {datetime.now().strftime('%d %B %Y')}")
    run.italic = True
    run.font.size = Pt(10)

    document.add_heading("Original Request", level=2)
    document.add_paragraph(request.strip())

    if tasks:
        document.add_heading("Agent Execution Plan", level=2)
        document.add_paragraph(
            "Before writing this document, the agent broke the request "
            "down into the following tasks:"
        )
        for task in tasks:
            p = document.add_paragraph(style="List Bullet")
            label = getattr(task, "title", str(task))
            run = p.add_run(label)
            run.bold = True
            description = getattr(task, "description", None)
            if description:
                p.add_run(f" — {description}")

    document.add_heading("Generated Content", level=2)
    for section_title, content in outputs.items():
        document.add_heading(section_title, level=3)
        _add_section_content(document, content)

    document.add_page_break()
    document.add_heading("Document Information", level=2)
    document.add_paragraph(
        f"Generated on: {datetime.now().strftime('%d %B %Y %H:%M')}"
    )

    filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = OUTPUT_DIR / filename
    document.save(filepath)

    return str(filepath)